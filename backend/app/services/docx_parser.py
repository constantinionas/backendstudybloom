import re
from pathlib import Path

from docx import Document

from app.schemas.upload_schema import AnswerPreview, QuestionPreview


QUESTION_PATTERN = re.compile(r"^\s*(\d+)[\.\)]\s+(.*\S)\s*$")
ANSWER_PATTERN = re.compile(r"^\s*([a-jA-J])[\.\)]\s+(.*\S)\s*$")


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def paragraph_has_highlight(paragraph) -> bool:
    """
    Verifică dacă minimum un fragment de text din paragraf este evidențiat.
    """
    for run in paragraph.runs:
        if run.font.highlight_color is not None:
            return True

    return False


def parse_highlighted_paragraphs(document: Document) -> list[QuestionPreview]:
    questions: list[QuestionPreview] = []
    current_question: dict | None = None

    for paragraph in document.paragraphs:
        text = normalize_text(paragraph.text)

        if not text:
            continue

        question_match = QUESTION_PATTERN.match(text)

        if question_match:
            if current_question is not None and current_question["answers"]:
                questions.append(
                    QuestionPreview(
                        number=current_question["number"],
                        text=current_question["text"],
                        answers=current_question["answers"],
                        source_pages=[],
                    )
                )

            current_question = {
                "number": int(question_match.group(1)),
                "text": question_match.group(2),
                "answers": [],
            }
            continue

        if current_question is None:
            continue

        answer_match = ANSWER_PATTERN.match(text)

        if answer_match:
            current_question["answers"].append(
                AnswerPreview(
                    label=answer_match.group(1).lower(),
                    text=answer_match.group(2),
                    correct=paragraph_has_highlight(paragraph),
                )
            )
            continue

        # Dacă nu este o nouă variantă, continuă ultimul răspuns găsit.
        if current_question["answers"]:
            last_answer = current_question["answers"][-1]
            last_answer.text = normalize_text(f"{last_answer.text} {text}")

            if paragraph_has_highlight(paragraph):
                last_answer.correct = True
        else:
            current_question["text"] = normalize_text(
                f'{current_question["text"]} {text}'
            )

    if current_question is not None and current_question["answers"]:
        questions.append(
            QuestionPreview(
                number=current_question["number"],
                text=current_question["text"],
                answers=current_question["answers"],
                source_pages=[],
            )
        )

    return questions


def parse_correct_wrong_tables(document: Document, start_number: int) -> list[QuestionPreview]:
    """
    Transformă fiecare tabel Corect / Greșit într-o întrebare cu variante multiple.
    Coloana cu valoarea 'Corect' marchează răspunsurile corecte.
    """
    questions: list[QuestionPreview] = []
    question_number = start_number

    for table in document.tables:
        rows = table.rows

        if len(rows) < 2:
            continue

        header_values = [
            normalize_text(cell.text).lower()
            for cell in rows[0].cells
        ]

        contains_correct_wrong_header = any(
            "corect" in value and ("gresit" in value or "greșit" in value)
            for value in header_values
        )

        if not contains_correct_wrong_header:
            continue

        cells_in_header = [normalize_text(cell.text) for cell in rows[0].cells]

        if len(cells_in_header) >= 2:
            question_text = cells_in_header[1]
        else:
            question_text = cells_in_header[0]

        answers: list[AnswerPreview] = []

        for row in rows[1:]:
            values = [normalize_text(cell.text) for cell in row.cells]

            if len(values) < 2:
                continue

            answer_label = values[0].replace(".", "").strip().lower()
            answer_text = values[1]

            verdict = values[-1].lower()
            is_correct = verdict.startswith("corect")

            if not answer_text:
                continue

            answers.append(
                AnswerPreview(
                    label=answer_label or chr(97 + len(answers)),
                    text=answer_text,
                    correct=is_correct,
                )
            )

        if answers:
            questions.append(
                QuestionPreview(
                    number=question_number,
                    text=question_text,
                    answers=answers,
                    source_pages=[],
                )
            )
            question_number += 1

    return questions


def parse_docx_questions(file_path: str | Path) -> tuple[list[QuestionPreview], list[str]]:
    document = Document(str(file_path))
    warnings: list[str] = []

    questions = parse_highlighted_paragraphs(document)

    next_number = max((question.number for question in questions), default=0) + 1
    table_questions = parse_correct_wrong_tables(document, next_number)

    questions.extend(table_questions)

    for question in questions:
        if not any(answer.correct for answer in question.answers):
            warnings.append(
                f"Întrebarea {question.number} nu are niciun răspuns corect detectat."
            )

    if not questions:
        warnings.append(
            "Nu am găsit întrebări compatibile în fișierul DOCX."
        )

    return questions, warnings